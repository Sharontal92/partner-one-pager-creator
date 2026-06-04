#!/usr/bin/env python3
"""
Partner One Pager — DOCX Generator
Produces a styled Word document following the Optimove partner one-pager template.

Usage:
    python generate_docx.py <partner_data.json> <output_dir>
"""

import sys
import json
import os
import re
import pathlib
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx --break-system-packages -q")
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Brand Colors (as RGB tuples) ───────────────────────────────────────────────
MIDNIGHT_VIOLET = RGBColor(0x30, 0x2c, 0x69)
LIME_GLOW       = RGBColor(0xDF, 0xF6, 0x70)
SHADOW_BLACK    = RGBColor(0x11, 0x11, 0x11)
CLOUD_TINT      = RGBColor(0xEF, 0xEF, 0xEF)
STORM_TINT      = RGBColor(0xD3, 0xD3, 0xD3)
SIGNAL_MIST     = RGBColor(0x9E, 0x97, 0xCB)
PURE_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Font paths ─────────────────────────────────────────────────────────────────
SKILL_DIR = Path(__file__).parent.parent
# Try several candidate paths for the font directory
_font_candidates = [
    SKILL_DIR.parent / "optimove-brand-bible" / "assets" / "fonts" / "Poppins",
    SKILL_DIR.parent.parent / "optimove-brand-bible" / "assets" / "fonts" / "Poppins",
    pathlib.Path("/sessions").glob("*/mnt/.claude/skills/optimove-brand-bible/assets/fonts/Poppins"),
]
FONT_DIR = next((p for p in _font_candidates if isinstance(p, pathlib.Path) and p.exists()), _font_candidates[0])
POPPINS   = "Poppins"  # DOCX references font by name; Poppins must be installed on viewer's machine
                       # or embedded. We set the name and embed the font.

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"),
                    color="D3D3D3", sz="4"):
    """Set borders on specific sides of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False,
            size_pt=11, color=None, font_name=POPPINS):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1, color=SHADOW_BLACK, size_pt=38,
                align=WD_ALIGN_PARAGRAPH.LEFT, space_after_mm=4):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Mm(space_after_mm)
    para.paragraph_format.space_before = Mm(0)
    add_run(para, text, bold=True, size_pt=size_pt, color=color)
    return para


def add_body(doc, text, size_pt=11, color=SHADOW_BLACK,
             space_after_mm=3, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Mm(space_after_mm)
    para.paragraph_format.space_before = Mm(0)
    add_run(para, text, size_pt=size_pt, color=color)
    return para


def add_divider(doc):
    """Add a thin horizontal rule."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Mm(2)
    para.paragraph_format.space_after  = Mm(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:color"), "D3D3D3")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)


# ── Header row ─────────────────────────────────────────────────────────────────
def add_header_table(doc, partner_name, content_width_mm=170):
    """Adds a 2-column header row: Optimove | Partner."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    col_w = Mm(content_width_mm / 2)
    table.columns[0].width = col_w
    table.columns[1].width = col_w

    left_cell  = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # Optimove label
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(lp, "OPTIMOVE", bold=True, size_pt=14, color=MIDNIGHT_VIOLET)

    # Partner name
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(rp, partner_name, bold=True, size_pt=12, color=SHADOW_BLACK)

    # Style the table — no internal borders, just a bottom line
    for cell in [left_cell, right_cell]:
        set_cell_border(cell, sides=("bottom",), color="D3D3D3", sz="4")
        set_cell_border(cell, sides=("top", "left", "right"), color="FFFFFF", sz="0")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Row height
    row = table.rows[0]
    row.height = Mm(14)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Mm(8)


# ── Page builders ──────────────────────────────────────────────────────────────
def build_page1(doc, data, content_width_mm=170):
    add_header_table(doc, data["partner_name"], content_width_mm)

    # Headline
    add_heading(doc, data["headline"], size_pt=36, space_after_mm=6)

    # Intro
    intro = data.get("intro", "")
    for para_text in [p.strip() for p in intro.split("\n\n") if p.strip()]:
        add_body(doc, para_text, space_after_mm=3)

    add_divider(doc)

    # Capabilities
    for i, cap in enumerate(data.get("capabilities", [])):
        # Bullet indicator using a styled table
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.columns[0].width = Mm(32)
        table.columns[1].width = Mm(content_width_mm - 32)

        # Left: colored "pill" block as a simple filled cell
        left = table.cell(0, 0)
        bg_colors = ["302c69", "dff670", "9e97cb"]
        fg_colors = [PURE_WHITE, MIDNIGHT_VIOLET, PURE_WHITE]
        set_cell_bg(left, bg_colors[i % len(bg_colors)])
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        idx_symbol = ["●", "■", "▶"]
        add_run(lp, idx_symbol[i % len(idx_symbol)],
                bold=True, size_pt=20, color=fg_colors[i % len(fg_colors)])
        left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Right: title + description
        right = table.cell(0, 1)
        rp = right.paragraphs[0]
        rp.paragraph_format.space_before = Mm(3)
        add_run(rp, cap["title"] + "\n", bold=True, size_pt=13, color=MIDNIGHT_VIOLET)
        rp2 = right.add_paragraph()
        rp2.paragraph_format.space_after = Mm(2)
        add_run(rp2, cap["description"], size_pt=10, color=SHADOW_BLACK)

        # Remove all table borders
        for cell in [left, right]:
            set_cell_border(cell, sides=("top","bottom","left","right"),
                            color="FFFFFF", sz="0")

        doc.add_paragraph().paragraph_format.space_after = Mm(3)


def build_page2(doc, data, content_width_mm=170):
    add_header_table(doc, data["partner_name"], content_width_mm)

    # Use cases
    add_heading(doc, "Seamlessly power any use case",
                size_pt=22, space_after_mm=5)

    for uc in data.get("use_cases", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Mm(1)
        add_run(p, uc["title"], bold=True, size_pt=12, color=MIDNIGHT_VIOLET)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Mm(4)
        add_run(p2, uc["description"], size_pt=10, color=SHADOW_BLACK)

    add_divider(doc)

    # Social proof
    trusted_by = data.get("trusted_by", "1,200+")
    add_heading(doc, f"Trusted by {trusted_by} Brands",
                size_pt=18, align=WD_ALIGN_PARAGRAPH.CENTER,
                color=SHADOW_BLACK, space_after_mm=4)

    clients = data.get("client_logos", [])
    if clients:
        logos = (clients + [""] * 6)[:6]
        rows = [logos[:3], logos[3:]]
        col_w = Mm(content_width_mm / 3)

        grid = doc.add_table(rows=2, cols=3)
        grid.alignment = WD_TABLE_ALIGNMENT.CENTER
        grid.autofit = False
        for col in grid.columns:
            col.width = col_w

        for r_i, row_data in enumerate(rows):
            for c_i, name in enumerate(row_data):
                cell = grid.cell(r_i, c_i)
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if name:
                    add_run(cp, name, bold=True, size_pt=10, color=SHADOW_BLACK)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                row = grid.rows[r_i]
                row.height = Mm(14)
                set_cell_border(cell, sides=("top","bottom","left","right"),
                                color="D3D3D3", sz="4")

    # Optional quote
    quote = data.get("client_quote")
    if quote and quote.get("text"):
        add_divider(doc)
        qp = doc.add_paragraph()
        qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qp.paragraph_format.left_indent  = Mm(15)
        qp.paragraph_format.right_indent = Mm(15)
        add_run(qp, f"“{quote['text']}”",
                italic=True, size_pt=11, color=MIDNIGHT_VIOLET)
        if quote.get("attribution"):
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(ap, quote["attribution"], bold=True, size_pt=9, color=SHADOW_BLACK)

    add_divider(doc)

    # KPIs
    cta = data.get("cta", "Do more when you go Positionless")
    add_heading(doc, cta, size_pt=18,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                color=SHADOW_BLACK, space_after_mm=4)

    kpis = data.get("kpis", [])
    if kpis:
        n = min(len(kpis), 3)
        col_w = Mm(content_width_mm / n)
        kt = doc.add_table(rows=2, cols=n)
        kt.alignment = WD_TABLE_ALIGNMENT.CENTER
        kt.autofit = False
        for col in kt.columns:
            col.width = col_w

        for i, kpi in enumerate(kpis[:n]):
            # Value row
            vc = kt.cell(0, i)
            vp = vc.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(vp, kpi["value"], bold=True, italic=True,
                    size_pt=32, color=MIDNIGHT_VIOLET)
            kt.rows[0].height = Mm(16)

            # Label row
            lc = kt.cell(1, i)
            lp2 = lc.paragraphs[0]
            lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(lp2, kpi["label"], size_pt=9, color=SHADOW_BLACK)
            kt.rows[1].height = Mm(10)

            # Borders
            for row_i in range(2):
                cell = kt.cell(row_i, i)
                borders = ["top", "bottom", "left", "right"]
                if i < n - 1:
                    set_cell_border(cell, sides=("right",), color="D3D3D3", sz="4")
                set_cell_border(cell, sides=("top", "bottom", "left"),
                                color="FFFFFF", sz="0")


# ── Main ───────────────────────────────────────────────────────────────────────
def generate(data_path: str, output_dir: str) -> str:
    with open(data_path) as f:
        data = json.load(f)

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    partner  = data["partner_name"]
    safe     = re.sub(r'[^\w\s-]', '', partner).strip()
    out_file = output_dir / f"Optimove x {safe} - One Pager.docx"

    doc = Document()

    # Page setup: A4, narrow margins
    section = doc.sections[0]
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(section, attr, Mm(20))

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = POPPINS
    style.font.size = Pt(11)
    style.font.color.rgb = SHADOW_BLACK

    build_page1(doc, data)
    add_page_break(doc)
    build_page2(doc, data)

    doc.save(str(out_file))
    print(f"✅ DOCX generated: {out_file}")
    return str(out_file)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: generate_docx.py <partner_data.json> <output_dir>")
        sys.exit(1)
    generate(sys.argv[1], sys.argv[2])
